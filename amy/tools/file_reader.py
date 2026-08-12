"""
tools/file_reader.py — Extract plain text from user-uploaded documents.

Supported formats:
  .txt / .md / .py / .js / .csv / .json / .html  → read directly as UTF-8
  .pdf                                             → pdfplumber (text layer)
  .docx                                            → python-docx
  .xlsx / .xls                                     → openpyxl / xlrd
  anything else                                    → try UTF-8, fallback error

All functions are sync (called via run_in_executor from brain).
"""
import io
import logging
import mimetypes

logger = logging.getLogger("amy.tools.file_reader")

# Max characters we pass to the LLM — keeps tokens sane
MAX_CHARS = 12_000
# Overlap notice appended when file is truncated
_TRUNCATION_NOTE = "\n\n[... file truncated to fit context ...]"


# ── Public API ────────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Main entry point.
    Returns extracted text (truncated if needed) or raises ValueError.
    """
    ext = _ext(filename)
    logger.info("Extracting text from '%s' (ext=%s, size=%d bytes)", filename, ext, len(file_bytes))

    try:
        if ext in (".txt", ".md", ".py", ".js", ".ts", ".css", ".html",
                   ".htm", ".json", ".yaml", ".yml", ".toml", ".ini",
                   ".env", ".sh", ".bat", ".log", ".xml", ".rst"):
            text = _read_plain(file_bytes)

        elif ext == ".pdf":
            text = _read_pdf(file_bytes)

        elif ext == ".docx":
            text = _read_docx(file_bytes)

        elif ext in (".xlsx", ".xls"):
            text = _read_excel(file_bytes, ext)

        elif ext == ".csv":
            text = _read_plain(file_bytes)   # CSV is plain text

        else:
            # Try plain text as last resort
            logger.warning("Unknown extension '%s', trying UTF-8 decode", ext)
            text = _read_plain(file_bytes)

    except ImportError as e:
        missing = str(e).split("'")[1] if "'" in str(e) else str(e)
        raise ValueError(
            f"Cannot read {ext} files — missing library: {missing}. "
            "Add it to requirements.txt."
        )
    except Exception as e:
        logger.error("Extraction failed for '%s': %s", filename, e)
        raise ValueError(f"Could not read this file: {e}")

    if not text or not text.strip():
        raise ValueError("The file appears to be empty or has no readable text.")

    # Truncate if needed
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + _TRUNCATION_NOTE
        logger.info("File truncated to %d chars", MAX_CHARS)

    return text


def describe_file(filename: str, text: str) -> str:
    """
    Returns a short meta-description to include in the LLM prompt.
    """
    lines = text.count("\n") + 1
    words = len(text.split())
    return (
        f"[File: {filename} | ~{words} words | ~{lines} lines]"
    )


# ── Format handlers ───────────────────────────────────────────────────────────

def _read_plain(file_bytes: bytes) -> str:
    """Decode bytes as UTF-8 with fallback to latin-1."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


def _read_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber")

    pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                pages.append(f"[Page {i + 1}]\n{page_text.strip()}")

    return "\n\n".join(pages)


def _read_docx(file_bytes: bytes) -> str:
    """Extract text from .docx using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx")

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _read_excel(file_bytes: bytes, ext: str) -> str:
    """Extract text from .xlsx / .xls using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl")

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    return "\n\n".join(sheets)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _ext(filename: str) -> str:
    """Return lowercase file extension including dot, e.g. '.pdf'"""
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()
